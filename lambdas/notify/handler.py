import io
import json
import os
import re
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

import boto3
from docx import Document
from docx.shared import Pt, Inches, RGBColor

REGION = os.environ.get("AWS_REGION", "ap-east-1")
s3_client = boto3.client(
    "s3",
    region_name=REGION,
    endpoint_url=f"https://s3.{REGION}.amazonaws.com",
    config=boto3.session.Config(signature_version="s3v4"),
)
dynamodb = boto3.resource("dynamodb")

JOBS_TABLE = os.environ["JOBS_TABLE"]
DATA_BUCKET = os.environ["DATA_BUCKET"]
SES_FROM_EMAIL = os.environ["SES_FROM_EMAIL"]

ses_client = boto3.client("sesv2", region_name=os.environ.get("SES_REGION", "ap-southeast-1"))


def lambda_handler(event, context):
    job_id = event["job_id"]
    email = event["email"]
    status = event.get("status", "completed")
    error = event.get("error")

    table = dynamodb.Table(JOBS_TABLE)
    job_record = table.get_item(Key={"job_id": job_id}).get("Item", {})
    file_name = job_record.get("file_name", "recording")

    if status == "completed":
        _send_success_email(job_id, email, file_name)
    else:
        _send_failure_email(job_id, email, error, file_name)

    return {"status": "notified", "email": email}


# ---------------------------------------------------------------------------
# Success email with DOCX attachment
# ---------------------------------------------------------------------------
def _send_success_email(job_id, email, file_name):
    key = f"jobs/{job_id}/meeting_minutes.md"
    resp = s3_client.get_object(Bucket=DATA_BUCKET, Key=key)
    md_content = resp["Body"].read().decode("utf-8")

    # Build DOCX
    docx_bytes = _md_to_docx(md_content, file_name)

    # Build display name from file_name (strip extension + timestamp)
    display_name = re.sub(r'-\d{8}_\d{6}-.*$', '', file_name)
    if not display_name:
        display_name = file_name.rsplit(".", 1)[0]
    docx_filename = f"{display_name} — 會議紀錄.docx"

    # Build MIME email with attachment
    msg = MIMEMultipart("mixed")
    msg["Subject"] = f"你嘅會議紀錄已準備好 — {display_name}"
    msg["From"] = SES_FROM_EMAIL
    msg["To"] = email

    # HTML body — short notification, no inline minutes
    html_body = f"""
    <html>
    <body style="font-family: -apple-system, 'Segoe UI', sans-serif; max-width: 600px; margin: 0 auto; padding: 20px;">
        <h2 style="color: #3a7f82;">會議紀錄已準備好</h2>
        <p style="font-size: 16px; color: #333;">
            <strong>{display_name}</strong> 嘅會議紀錄已經處理完成。
        </p>
        <p style="color: #666;">
            請查收附件中嘅 Word 文件（.docx）。
        </p>
        <hr style="border: none; border-top: 1px solid #e0e0e0; margin: 24px 0;">
        <p style="font-size: 12px; color: #999;">
            此郵件由 Precis（精記）自動發送。
        </p>
    </body>
    </html>
    """
    text_body = f"{display_name} 嘅會議紀錄已經處理完成。請查收附件中嘅 Word 文件。"

    # Attach HTML + text body
    body_part = MIMEMultipart("alternative")
    body_part.attach(MIMEText(text_body, "plain", "utf-8"))
    body_part.attach(MIMEText(html_body, "html", "utf-8"))
    msg.attach(body_part)

    # Attach DOCX
    attachment = MIMEApplication(docx_bytes, _subtype="vnd.openxmlformats-officedocument.wordprocessingml.document")
    attachment.add_header("Content-Disposition", "attachment", filename=("utf-8", "", docx_filename))
    msg.attach(attachment)

    send_params = {
        "FromEmailAddress": SES_FROM_EMAIL,
        "Destination": {"ToAddresses": [email]},
        "Content": {"Raw": {"Data": msg.as_bytes()}},
    }
    ses_client.send_email(**send_params)


# ---------------------------------------------------------------------------
# Markdown → DOCX conversion
# ---------------------------------------------------------------------------
def _md_to_docx(md_content, file_name):
    """Convert meeting minutes Markdown to a formatted Word document."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Style defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.color.rgb = RGBColor(0x3a, 0x7f, 0x82)

    lines = md_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)

        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            # Add a thin line via bottom border
            run = p.add_run("─" * 60)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # Bullet list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            text = _clean_md_inline(text)
            doc.add_paragraph(text, style="List Bullet")

        # Numbered list
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            text = _clean_md_inline(text)
            doc.add_paragraph(text, style="List Number")

        # Table (simple markdown table)
        elif stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, table_lines)
            continue  # skip i += 1 at bottom

        # Regular paragraph
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, stripped)

        i += 1

    # Save to bytes in memory
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _clean_md_inline(text):
    """Remove markdown inline formatting for plain text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def _add_formatted_runs(paragraph, text):
    """Add runs with bold/italic formatting from markdown inline syntax."""
    # Split on **bold** and *italic* patterns
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x44, 0x44)
        else:
            paragraph.add_run(part)


def _add_table(doc, table_lines):
    """Parse markdown table lines and add a Word table."""
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        # Skip separator rows (---)
        if all(re.match(r'^[-:]+$', c) for c in cells):
            continue
        rows.append(cells)

    if len(rows) < 1:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Light Grid Accent 1"

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = _clean_md_inline(cell_text)
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True


# ---------------------------------------------------------------------------
# Failure email (simple, no attachment)
# ---------------------------------------------------------------------------
def _extract_error_reason(error):
    if not error:
        return "未知錯誤"
    try:
        cause = error.get("Cause", "")
        if isinstance(cause, str):
            try:
                cause_obj = json.loads(cause)
            except (json.JSONDecodeError, TypeError):
                cause_obj = {}
        else:
            cause_obj = cause

        reason = ""
        container = cause_obj.get("Container", {})
        if container.get("Reason"):
            reason = container["Reason"]
        elif cause_obj.get("StatusReason"):
            reason = cause_obj["StatusReason"]
        elif cause_obj.get("errorMessage"):
            reason = cause_obj["errorMessage"]

        if reason:
            if len(reason) > 300:
                reason = reason[:300] + "..."
            return reason

        error_type = error.get("Error", "")
        if error_type:
            return f"錯誤類型：{error_type}"

        return "未知錯誤"
    except Exception:
        return "處理錯誤信息時發生異常"


def _send_failure_email(job_id, email, error, file_name):
    error_msg = _extract_error_reason(error)

    send_params = {
        "FromEmailAddress": SES_FROM_EMAIL,
        "Destination": {"ToAddresses": [email]},
        "Content": {
            "Simple": {
                "Subject": {"Data": f"會議紀錄處理失敗 — {file_name}", "Charset": "UTF-8"},
                "Body": {
                    "Text": {
                        "Data": f"「{file_name}」處理失敗。\n\n"
                        f"原因：{error_msg}\n\n"
                        "請嘗試重新上傳錄影。如果問題持續，請聯絡管理員。",
                        "Charset": "UTF-8",
                    },
                },
            },
        },
    }
    ses_client.send_email(**send_params)
